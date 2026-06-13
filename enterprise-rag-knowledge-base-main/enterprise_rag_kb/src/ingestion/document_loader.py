"""
Enterprise RAG - Document Loader
==================================
Supports: PDF, TXT, Markdown, CSV, JSON, DOCX
Each loaded document becomes a Document object with content + metadata.
"""

import os
import csv
import json
import hashlib
import sys
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Optional, Dict, Any

sys.path.insert(0, str(Path(__file__).parent.parent.parent))
from src.logger import get_logger

logger = get_logger(__name__)


@dataclass
class Document:
    """Represents a loaded document with content and metadata."""
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    doc_id: str = ""

    def __post_init__(self):
        if not self.doc_id:
            self.doc_id = hashlib.md5(
                (self.content[:200] + str(self.metadata)).encode()
            ).hexdigest()[:12]

    def __repr__(self):
        preview = self.content[:80].replace("\n", " ")
        return f"Document(id={self.doc_id}, source={self.metadata.get('source','?')}, preview='{preview}...')"


class DocumentLoader:
    """
    Loads documents from files or directories.
    Supports multiple formats with automatic type detection.
    """

    SUPPORTED = {".pdf", ".txt", ".md", ".csv", ".json", ".docx"}

    def load_file(self, file_path: str) -> List[Document]:
        """Load a single file, auto-detecting format."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED:
            logger.warning(f"Unsupported format: {ext} - skipping {path.name}")
            return []

        logger.info(f"Loading [{ext.upper()}]: {path.name}")
        try:
            if ext == ".pdf":
                return self._load_pdf(path)
            elif ext in (".txt", ".md"):
                return self._load_text(path)
            elif ext == ".csv":
                return self._load_csv(path)
            elif ext == ".json":
                return self._load_json(path)
            elif ext == ".docx":
                return self._load_docx(path)
        except Exception as e:
            logger.error(f"Failed to load {path.name}: {e}")
            return []

    def load_directory(self, dir_path: str, recursive: bool = True) -> List[Document]:
        """Load all supported documents from a directory."""
        path = Path(dir_path)
        if not path.exists():
            raise FileNotFoundError(f"Directory not found: {dir_path}")

        pattern = "**/*" if recursive else "*"
        files = [f for f in path.glob(pattern)
                 if f.is_file() and f.suffix.lower() in self.SUPPORTED]

        logger.info(f"Found {len(files)} supported files in {dir_path}")
        all_docs = []
        for f in files:
            docs = self.load_file(str(f))
            all_docs.extend(docs)

        logger.info(f"Loaded {len(all_docs)} documents total")
        return all_docs

    # ── Format-Specific Loaders ─────────────────────────────────────────────

    def _load_pdf(self, path: Path) -> List[Document]:
        try:
            import pypdf
        except ImportError:
            logger.warning("pypdf not installed. Install with: pip install pypdf")
            return self._load_text_fallback(path)

        docs = []
        with open(path, "rb") as f:
            reader = pypdf.PdfReader(f)
            total_pages = len(reader.pages)
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                text = text.strip()
                if len(text) < 20:
                    continue
                docs.append(Document(
                    content=text,
                    metadata={
                        "source": str(path),
                        "filename": path.name,
                        "file_type": "pdf",
                        "page": i + 1,
                        "total_pages": total_pages,
                    }
                ))
        return docs

    def _load_text(self, path: Path) -> List[Document]:
        encodings = ["utf-8", "latin-1", "cp1252"]
        for enc in encodings:
            try:
                content = path.read_text(encoding=enc)
                return [Document(
                    content=content.strip(),
                    metadata={
                        "source": str(path),
                        "filename": path.name,
                        "file_type": path.suffix.lstrip("."),
                        "encoding": enc,
                        "size_bytes": path.stat().st_size,
                    }
                )]
            except UnicodeDecodeError:
                continue
        return []

    def _load_csv(self, path: Path) -> List[Document]:
        docs = []
        with open(path, "r", encoding="utf-8-sig", errors="replace") as f:
            reader = csv.DictReader(f)
            rows = list(reader)

        if not rows:
            return []

        # Convert entire CSV to readable text
        columns = list(rows[0].keys())
        text_parts = [f"CSV Document: {path.name}", f"Columns: {', '.join(columns)}", "---"]
        for i, row in enumerate(rows):
            row_text = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
            text_parts.append(row_text)

        full_text = "\n".join(text_parts)

        # Also create individual row documents for large CSVs
        if len(rows) > 100:
            batch_size = 50
            for i in range(0, len(rows), batch_size):
                batch = rows[i:i+batch_size]
                batch_text = "\n".join(
                    " | ".join(f"{k}: {v}" for k, v in row.items() if v)
                    for row in batch
                )
                docs.append(Document(
                    content=batch_text,
                    metadata={
                        "source": str(path),
                        "filename": path.name,
                        "file_type": "csv",
                        "row_start": i + 1,
                        "row_end": i + len(batch),
                    }
                ))
        else:
            docs.append(Document(
                content=full_text,
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "file_type": "csv",
                    "num_rows": len(rows),
                    "columns": columns,
                }
            ))
        return docs

    def _load_json(self, path: Path) -> List[Document]:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)

        # Handle list of records
        if isinstance(data, list):
            docs = []
            for i, item in enumerate(data):
                text = json.dumps(item, indent=2, ensure_ascii=False)
                docs.append(Document(
                    content=text,
                    metadata={
                        "source": str(path),
                        "filename": path.name,
                        "file_type": "json",
                        "record_index": i,
                    }
                ))
            return docs
        else:
            return [Document(
                content=json.dumps(data, indent=2, ensure_ascii=False),
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "file_type": "json",
                }
            )]

    def _load_docx(self, path: Path) -> List[Document]:
        try:
            import docx
        except ImportError:
            logger.warning("python-docx not installed. Install with: pip install python-docx")
            return []

        doc = docx.Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)

        return [Document(
            content=content,
            metadata={
                "source": str(path),
                "filename": path.name,
                "file_type": "docx",
                "num_paragraphs": len(paragraphs),
            }
        )]

    def _load_text_fallback(self, path: Path) -> List[Document]:
        try:
            return self._load_text(path)
        except Exception:
            return []
